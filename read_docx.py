import docx
from docx import Document
import re
import use_to_wave
import pyttsx3
from pydub import AudioSegment  # 先导入这个模块
import os


'''
先通过pyttsx3发声
然后通过特殊手段录制系统声
outfile = "out.aiff"
    tts = pyttsx3.init()
    tts.save_to_file(content, outfile)
    tts.runAndWait()
'''


'''
add_mp3
将两个音频合并的程序
输入第一个音频的地址str和第二个音频的地址str
输出到第一个音频的地址
'''
# -*- coding: utf-8 -*-
def add_mp3(music1,music2):
    # 加载需要合并的两个mp3音频
    input_music_1 = AudioSegment.from_mp3(music1)
    input_music_2 = AudioSegment.from_mp3(music2)
    #获取两个音频的响度（音量）
    input_music_1_db = input_music_1.dBFS
    input_music_2_db = input_music_2.dBFS
    # 获取两个音频的时长，单位为毫秒
    input_music_1_time = len(input_music_1)
    input_music_2_time = len(input_music_2)
    # 调整两个音频的响度一致
    '''
    db = input_music_1_db- input_music_2_db
    if db > 0:
        input_music_1 += abs(db)
    elif db < 0:
        input_music_2 += abs(db)
    '''
    # 合并音频
    output_music = input_music_1 + input_music_2
    # 简单输入合并之后的音频
    output_music.export(music1, format="wav")# 前面是保存路径，后面是保存格式
    #复杂输入合并之后的音频
    # bitrate：比特率，album：专辑名称，artist：歌手，cover：封面图片
    #output_music.export("E:/output_music.mp3", format="mp3", bitrate="192k", tags={"album": "专辑", "artist": "歌手"}, cover="E:/封面.jpg")
    print(len(output_music), output_music.channels)# 合并音频的时长，音频的声道，1是单声道，2是立体声




'''
生成英文字符以后将英文字符的有声小说转换为中文字符的有声小说的小程序
中文名不能被qq视频识别所以暂时不做。
'''
def change_name(types):
    if types == 'mp3':
        a = open('./tlj.mp3','rb')
        b = open('./天国记.mp3','wb')
        b.write(a.read())
        a.close()
        #os.remove(r'./tlj.mp3')
        b.flush()
        b.close()
    if types == 'wav':
        a = open('./tlj.wav','rb')
        b = open('./天国记.wav', 'wb')
        b.write(a.read())
        a.close()
        #os.remove(r'./tlj.wav')
        b.flush()
        b.close()

'''
get_the_mp3()
输入types 对应wav或mp3
wav格式比较大，但是方便处理
MP3格式比较小。
'''
def word_to_voice(the_str):
    print(the_str)
    pattern = r'\.|/|;|\'|`|\[|\]|<|>|\?|"|\{|\}|\~|!|@|#|\$|%|\^|&|\(|\)|-|=|\_|\+|。|、|：|；|‘|’|【|】|·|！|…|（|）'
    result_list = re.split(pattern, the_str)
    for words in result_list:
        #print(words)
        engine = pyttsx3.init()
        #engine.say(words)
        engine.save_to_file(words,r'.\tample.wav')
        engine.runAndWait()
        add_mp3(r'./tlj.wav',r'.\tample.wav')
        #base64.b64encode(sound_wav_rb)
    os.remove(r'./tample.wav')
    

def get_the_mp3(types):
    doc = docx.Document('./天国记.docx')
    if types == 'wav':
        '''
        pyttsx3初始化需要一个aiff，不知道为什么
        '''
        outfile = "out.aiff"
        tts = pyttsx3.init()
        tts.save_to_file('就这', outfile)
        # engine.say(words)
        tts.save_to_file('e', r'.\tlj.wav')
        tts.runAndWait()
        '''
        正式开始阅读了。
        '''
        for para in doc.paragraphs:
            word_to_voice(para.text)
    if types == 'mp3':
        for para in doc.paragraphs:
            file = open('./天国记.txt','a+')
            file.write(str(para.text))
            file.close()
        use_to_wave.use_by_txt('./天国记.txt')
        os.remove(r'./天国记.txt')
    change_name(types)
#！！！！测试用，测试前请注释掉main
#get_the_mp3()