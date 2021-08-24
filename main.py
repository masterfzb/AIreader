import randomname
import save_dict
import data_base
import count_tools
import time
import NLP_eng
import docx
from docx import Document
import read_docx
import random
'''
    data_base中存放种族，位阶，姓名等数据
    create_things使用data_base中的数据创造名字及设定
    save_data用于保存数据
'''
savepath = './save/'
'''
添加特定数量的自定义种群随机角色。
不能对这类角色指定名字，只能指定种群。

'''
#增加
def add_data_life(life_list):
    #创建一个随机角色
    for items in range(0,100):
        #time.sleep(0.125)
        life_list.update(randomname.create_a_life(randomname.create_a_random_type(),"随机",life_list))
    #创建数个特定的种群
    num = input('请输入你要生成的特定种群个数：')
    if num == '':
        num = '50'
    type_list = randomname.create_a_definite_type()
    for life in range(0,eval(num)):
        #time.sleep(0.125)
        life_list.update(randomname.create_a_life(type_list,"随机",life_list))
    #创建一个特定的角色
    life_list.update(randomname.create_a_life(randomname.create_a_definite_type(),"自定义",life_list))
    data_base.show_dic(life_list)
    return life_list
'''
创建主角的函数：
输入已创建生物的life_list
假如没有这个人，则创建一个角色作为the_one 并将他并入生物的life_list
输出决定了主角后的life_list和主角的the_one，类型都是dict

'''
def choose_the_one(life_list):
    print("游戏开始，以下人物为可供选择的人物。想必你会选择已经指定了名字的角色吧。")
    data_base.show_data(life_list.keys())
    name = input("游戏开始，请选择一个角色作为你的第一人称主角。")
    if name in life_list.keys():
        the_one = {name:life_list[name]}
        return the_one,life_list,name
    else:
        print("我们的数据库里没有这个人，但是您可以自己创建这个人。")
        the_one = randomname.create_a_life(randomname.create_a_definite_type(),"自定义",life_list)
        life_list.update(the_one)
        for items in the_one.keys():
            items = items
        return the_one,life_list,items
def main():
    try:
        doc1 = docx.Document('./天国记.docx')
    except:
        doc1 = docx.Document()
    str_to_save = ""
    life_list = dict()
    try:
        life_list = save_dict.read_a_dic(savepath,'life_list')
    except:
        print("未读取到新的数据！")
    #生成角色数据
    count_tools.show_dic_data(life_list)#展示dic的种类参数
    life_list = add_data_life(life_list)#运行添加人物的程序
    count_tools.show_dic_data(life_list)
    save_dict.save_a_dic(life_list,savepath,"life_list")#将新生成的人物存入库
    #data_base.show_dic(life_list)
    the_one,life_list,name = choose_the_one(life_list)
    data_base.show_dic(the_one)
    str_to_save = '你的名字是' + data_base.get_all_in_lists(data_base.geting_str(the_one.keys()) + '，看过你的人都认为:\n')
    #
    str_to_save += NLP_eng.nlp_eng_discribing_things(name,life_list,'生物')
    for items in range(0,3):
        name2 = random.choice(list(life_list.keys()))
        str_to_save += '\n迎面走来的是' +name2 + '\n在你的眼里，他是这样的一个：' + life_list[name2]['种族'] + ' :\n'
        str_to_save += NLP_eng.nlp_eng_discribing_things(name2,life_list,'生物')
        str_to_save += '你向'+name2+'打了个招呼，他一见你便大喊：“哇，美'+the_one[name]['性别']+'！”'

    doc1.add_paragraph(str_to_save)
    doc1.save('./天国记.docx')

    read_docx.get_the_mp3('wav')#将天国记读出来。
    read_docx.get_the_mp3('mp3')
    


if __name__ == '__main__':
    main()
